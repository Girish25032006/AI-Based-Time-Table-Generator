from flask import Blueprint, request, jsonify, send_from_directory
import mysql.connector

import os
import re
import pdfplumber
import pytesseract

from PIL import Image
from werkzeug.utils import secure_filename

from docx import Document
from openpyxl import load_workbook


faculty_api = Blueprint("faculty_api", __name__)


# ============================================================
# FACULTY IMPORT CONFIGURATION
# ============================================================

FACULTY_UPLOAD_FOLDER = "uploads/faculty"

os.makedirs(
    FACULTY_UPLOAD_FOLDER,
    exist_ok=True
)

ALLOWED_FACULTY_EXTENSIONS = {
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".docx",
    ".xlsx"
}


# ============================================================
# DATABASE CONNECTION
# ============================================================

def get_connection():

    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="root",
        database="timetable_db"
    )


# ============================================================
# WORKLOAD CALCULATION
# ============================================================

def get_workload_for_designation(designation):

    if not designation:
        return 0

    value = designation.lower().strip()

    # Principal → 6 hours
    if "principal" in value:
        return 6

    # HOD / Head → 12 hours
    if "hod" in value or "head" in value:
        return 12

    # Assistant Professor → 18 hours
    if "assistant professor" in value:
        return 18

    # Associate Professor → 16 hours
    if "associate professor" in value:
        return 16

    # Professor → 16 hours
    if "professor" in value:
        return 16

    return 0


# ============================================================
# BASIC TEXT CLEANING
# ============================================================

def clean_import_text(value):

    if value is None:
        return ""

    return re.sub(
        r"\s+",
        " ",
        str(value)
    ).strip()


# ============================================================
# PDF TEXT EXTRACTION
# ============================================================

def extract_pdf_text(file_path):

    text = ""

    try:

        with pdfplumber.open(file_path) as pdf:

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:

                    text += "\n" + page_text

    except Exception as error:

        print(
            "PDF extraction error:",
            error
        )

    return text


# ============================================================
# IMAGE OCR
# ============================================================

def extract_image_text(file_path):

    try:

        image = Image.open(
            file_path
        )

        text = pytesseract.image_to_string(
            image
        )

        return text

    except Exception as error:

        print(
            "Image OCR error:",
            error
        )

        return ""


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(file_path):

    text = ""

    try:

        document = Document(
            file_path
        )

        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                text += (
                    "\n" +
                    paragraph.text
                )

        for table in document.tables:

            for row in table.rows:

                values = []

                for cell in row.cells:

                    values.append(
                        cell.text
                    )

                text += (
                    "\n" +
                    " | ".join(values)
                )

    except Exception as error:

        print(
            "DOCX extraction error:",
            error
        )

    return text


# ============================================================
# XLSX EXTRACTION
# ============================================================

def extract_xlsx_text(file_path):

    text = ""

    try:

        workbook = load_workbook(
            file_path,
            data_only=True
        )

        for sheet in workbook.worksheets:

            for row in sheet.iter_rows(
                values_only=True
            ):

                values = []

                for value in row:

                    if value is not None:

                        values.append(
                            str(value)
                        )

                if values:

                    text += (
                        "\n" +
                        " | ".join(values)
                    )

    except Exception as error:

        print(
            "XLSX extraction error:",
            error
        )

    return text


# ============================================================
# GENERAL FILE TEXT EXTRACTION
# ============================================================

def extract_faculty_text(
    file_path,
    extension
):

    extension = extension.lower()

    if extension == ".pdf":

        return extract_pdf_text(
            file_path
        )

    if extension in (
        ".png",
        ".jpg",
        ".jpeg",
        ".webp"
    ):

        return extract_image_text(
            file_path
        )

    if extension == ".docx":

        return extract_docx_text(
            file_path
        )

    if extension == ".xlsx":

        return extract_xlsx_text(
            file_path
        )

    return ""


# ============================================================
# ============================================================
# FACULTY PDF / DOCUMENT PARSER
# ============================================================

def normalize_line(value):
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def get_clean_lines(text):
    return [
        normalize_line(line)
        for line in (text or "").splitlines()
        if normalize_line(line)
    ]


def extract_value_from_same_line(lines, labels):
    """
    Extract values such as:
        Designation: Assistant Professor
        Teaching Experience 25 Years
        Date of Joining the Institution 15-04-2024
    """
    for line in lines:
        lower_line = line.lower()

        for label in labels:
            label_lower = label.lower()

            if not lower_line.startswith(label_lower):
                continue

            remaining = line[len(label):].strip(" :|-")

            if remaining:
                return normalize_line(remaining)

    return ""


def extract_section(lines, start_keywords, stop_keywords):
    """
    Collect every line after a section heading until the next
    known section heading.
    """
    start_index = -1
    start_set = {x.lower() for x in start_keywords}
    stop_set = {x.lower() for x in stop_keywords}

    for index, line in enumerate(lines):
        if line.lower().strip() in start_set:
            start_index = index
            break

    if start_index == -1:
        return ""

    collected = []

    for line in lines[start_index + 1:]:
        if line.lower().strip() in stop_set:
            break
        collected.append(line)

    return "\n".join(collected).strip()


def extract_labeled_block(
    lines,
    start_patterns,
    stop_patterns
):
    """
    Handles headings and headings-with-value, including wrapped
    multi-line content.
    """
    start_index = -1
    first_value = ""

    for index, line in enumerate(lines):
        lower_line = line.lower()

        for pattern in start_patterns:
            match = re.match(
                pattern,
                lower_line,
                re.IGNORECASE
            )

            if match:
                start_index = index

                # Capture text after the heading when present.
                original_match = re.match(
                    pattern,
                    line,
                    re.IGNORECASE
                )

                if original_match and original_match.lastindex:
                    first_value = normalize_line(
                        original_match.group(
                            original_match.lastindex
                        )
                    )

                break

        if start_index != -1:
            break

    if start_index == -1:
        return ""

    values = []

    if first_value:
        values.append(first_value)

    for line in lines[start_index + 1:]:
        lower_line = line.lower()

        should_stop = False

        for stop_pattern in stop_patterns:
            if re.match(
                stop_pattern,
                lower_line,
                re.IGNORECASE
            ):
                should_stop = True
                break

        if should_stop:
            break

        values.append(line)

    return normalize_line(" ".join(values))


def convert_join_date(value):
    if not value:
        return None

    value = normalize_line(value)

    match = re.search(
        r"\b(\d{1,2})[-/.](\d{1,2})[-/.](\d{2,4})\b",
        value
    )

    if match:
        day = match.group(1)
        month = match.group(2)
        year = match.group(3)

        if len(year) == 2:
            year = "20" + year

        return (
            f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        )

    from datetime import datetime

    for date_format in (
        "%d-%m-%Y",
        "%d/%m/%Y",
        "%d.%m.%Y",
        "%Y-%m-%d",
        "%Y/%m/%d",
        "%d-%m-%y",
        "%d/%m/%y"
    ):
        try:
            return datetime.strptime(
                value,
                date_format
            ).strftime("%Y-%m-%d")
        except ValueError:
            pass

    return None


def detect_faculty_name(text):
    lines = get_clean_lines(text)

    # ========================================================
    # PRIMARY SKIT FORMAT
    # ========================================================
    # FACULTY PROFILE
    # Mr. Asghar Pasha
    # Assistant Professor | AIML
    #
    # Also supports:
    # Prof.Sreenivasan A
    # Prof. Sreenivasan A
    # ========================================================

    for index, line in enumerate(lines):

        if line.lower().strip() == "faculty profile":

            # Look only at the next few lines.
            # This prevents the description from being detected
            # as the faculty name.
            for candidate in lines[index + 1:index + 5]:

                candidate = normalize_line(candidate)

                if not candidate:
                    continue

                # Must start with a faculty title.
                match = re.match(
                    r"^(Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.)(.*)$",
                    candidate,
                    re.IGNORECASE
                )

                if not match:
                    continue

                title = match.group(1)
                name_part = match.group(2).strip()

                # Remove accidental space after title.
                name_part = name_part.strip()

                # ------------------------------------------------
                # IMPORTANT:
                # Do not accept a description paragraph as name.
                # ------------------------------------------------

                if not name_part:
                    continue

                # Name should NOT contain sentence-like text.
                if any(
                    phrase in name_part.lower()
                    for phrase in [
                        "has ",
                        "is ",
                        "with ",
                        "years ",
                        "experience",
                        "secured ",
                        "specialized ",
                        "member ",
                        "assistant professor",
                        "associate professor"
                    ]
                ):
                    continue

                # Name should not contain too many words.
                if len(name_part.split()) > 8:
                    continue

                # Must contain at least one alphabetic character.
                if not re.search(
                    r"[A-Za-z]",
                    name_part
                ):
                    continue

                # Preserve original title + name.
                return (
                    title + name_part
                ).strip()

    # ========================================================
    # EXPLICIT LABEL
    # ========================================================

    value = extract_value_from_same_line(
        lines,
        [
            "faculty name",
            "name of faculty"
        ]
    )

    if value:
        return value

    # ========================================================
    # FALLBACK
    # ========================================================

    pattern = re.compile(
        r"^(Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.?)\s*"
        r"[A-Za-z][A-Za-z .'-]{1,80}$",
        re.IGNORECASE
    )

    for line in lines:

        line = normalize_line(line)

        if not pattern.match(line):
            continue

        lower_line = line.lower()

        # Never accept obvious description text.
        if any(
            phrase in lower_line
            for phrase in [
                " has ",
                " is ",
                " with ",
                " years ",
                " experience",
                " secured ",
                " specialized ",
                " member ",
                " teaching "
            ]
        ):
            continue

        return line

    return ""


def detect_designation(text):
    lines = get_clean_lines(text)

    value = extract_value_from_same_line(
        lines,
        [
            "designation",
            "position"
        ]
    )

    if value:
        return value

    # Actual SKIT format:
    # Assistant Professor | Artificial Intelligence and Machine Learning
    for line in lines:
        if "|" not in line:
            continue

        left = normalize_line(
            line.split("|", 1)[0]
        )

        lower_left = left.lower()

        if (
            "assistant professor" in lower_left
            or "associate professor" in lower_left
            or re.fullmatch(r"professor(?:\s*&\s*head.*)?", left, re.I)
            or "hod" in lower_left
        ):
            return left

    return ""


def detect_department(text, current_department=""):

    department_value = normalize_line(
        current_department
    ).lower()

    lower_text = (text or "").lower()

    # ========================================================
    # CHECK CURRENT / DIRECT DEPARTMENT VALUE FIRST
    # ========================================================

    # AIML
    if (
        "artificial intelligence and machine learning"
        in department_value
        or "artificial intelligence & machine learning"
        in department_value
        or department_value == "aiml"
    ):
        return "AIML"

    # CSE
    if (
        "computer science and engineering"
        in department_value
        or department_value == "cse"
    ):
        return "CSE"

    # ISE
    if (
        "information science and engineering"
        in department_value
        or department_value == "ise"
    ):
        return "ISE"

    # ECE
    if (
        "electronics and communication engineering"
        in department_value
        or department_value == "ece"
    ):
        return "ECE"

    # VLSI
    if (
        "vlsi design and technology"
        in department_value
        or "electronics engineering (vlsi"
        in department_value
        or department_value == "vlsi"
    ):
        return "VLSI"

    # ME
    if (
        "mechanical engineering"
        in department_value
        or department_value in ("mechanical", "me")
    ):
        return "ME"

    # CIVIL
    if (
        "civil engineering"
        in department_value
        or department_value in ("civil", "civ")
    ):
        return "CIV"

    # ========================================================
    # SCIENCES AND HUMANITIES
    # IMPORTANT: CHECK THIS BEFORE SEARCHING THE FULL PDF
    # ========================================================

    if (
        "sciences and humanities"
        in department_value
        or department_value in (
            "sh",
            "science and humanities",
            "sciences & humanities",
            "science & humanities"
        )
    ):
        return "SH"

    # ========================================================
    # NOW SEARCH THE COMPLETE PDF TEXT
    # ========================================================

    # IMPORTANT:
    # SH must be checked BEFORE VLSI.
    #
    # The SKIT footer contains all departments, including VLSI.
    # Therefore footer text must not incorrectly decide the
    # faculty's department when the actual department is SH.
    # ========================================================

    # --------------------------------------------------------
    # AIML
    # --------------------------------------------------------

    if (
        "assistant professor | artificial intelligence and machine learning"
        in lower_text
        or "artificial intelligence and machine learning"
        in lower_text
    ):
        return "AIML"

    # --------------------------------------------------------
    # SCIENCES AND HUMANITIES
    # --------------------------------------------------------

    if (
        "assistant professor | sciences and humanities"
        in lower_text
        or "sciences and humanities"
        in lower_text
    ):
        return "SH"

    # --------------------------------------------------------
    # CSE
    # --------------------------------------------------------

    if (
        "assistant professor | computer science and engineering"
        in lower_text
        or "computer science and engineering"
        in lower_text
    ):
        return "CSE"

    # --------------------------------------------------------
    # ISE
    # --------------------------------------------------------

    if (
        "assistant professor | information science and engineering"
        in lower_text
        or "information science and engineering"
        in lower_text
    ):
        return "ISE"

    # --------------------------------------------------------
    # ECE
    # --------------------------------------------------------

    if (
        "assistant professor | electronics and communication engineering"
        in lower_text
        or "electronics and communication engineering"
        in lower_text
    ):
        return "ECE"

    # --------------------------------------------------------
    # VLSI
    # --------------------------------------------------------

    if (
        "assistant professor | electronics engineering (vlsi"
        in lower_text
        or "vlsi design and technology"
        in lower_text
    ):
        return "VLSI"

    # --------------------------------------------------------
    # ME
    # --------------------------------------------------------

    if (
        "assistant professor | mechanical engineering"
        in lower_text
        or "professor | mechanical engineering"
        in lower_text
        or "mechanical engineering"
        in lower_text
    ):
        return "ME"

    # --------------------------------------------------------
    # CIVIL
    # --------------------------------------------------------

    if (
        "assistant professor | civil engineering"
        in lower_text
        or "professor | civil engineering"
        in lower_text
        or "civil engineering"
        in lower_text
    ):
        return "CIV"

    return ""


def parse_faculty_information(text):
    """
    Parser designed for the SKIT-style faculty profile PDF.

    It keeps complete multi-line sections instead of taking only
    the first line.
    """
    lines = get_clean_lines(text)

    # ========================================================
    # BASIC INFORMATION
    # ========================================================

    faculty_name = detect_faculty_name(text)
    designation = detect_designation(text)

    # Department is normally on the same line as designation.
    department = ""

    for line in lines:
        if "|" not in line:
            continue

        parts = line.split("|", 1)

        if len(parts) != 2:
            continue

        left = parts[0].lower()

        if (
            "assistant professor" in left
            or "associate professor" in left
            or "professor" in left
            or "hod" in left
        ):
            department = detect_department(
                parts[1]
            )

            if department:
                break

    if not department:
        department = detect_department(text)

    # ========================================================
    # DESCRIPTION
    # ========================================================

    # ========================================================
    # FACULTY DESCRIPTION
    # ========================================================

    faculty_description = ""

    # Find the designation line first.
    designation_line_index = -1

    for index, line in enumerate(lines):

        line = normalize_line(line)
        lower_line = line.lower()

        if "|" not in line:
            continue

        left_side = normalize_line(
            line.split("|", 1)[0]
        )

        left_lower = left_side.lower()

        if (
                "assistant professor" in left_lower
                or "associate professor" in left_lower
                or "professor" in left_lower
                or "head of the department" in left_lower
                or "hod" in left_lower
        ):
            designation_line_index = index
            break

    # Description is directly after designation.
    if designation_line_index != -1:

        description_parts = []

        for j in range(
                designation_line_index + 1,
                len(lines)
        ):

            next_line = normalize_line(
                lines[j]
            )

            if not next_line:
                continue

            next_lower = next_line.lower()

            # Stop at joining date.
            if (
                    "date of joining" in next_lower
                    or "joining date" in next_lower
            ):
                break

            # Stop at the next profile section.
            if next_lower in [
                "qualifications",
                "teaching experience",
                "skills",
                "research & publications",
                "research and publications"
            ]:
                break

            # Avoid footer text if PDF extraction reaches it.
            if (
                    "lets possibilities together" in next_lower
                    or "made by continum" in next_lower
            ):
                break

            description_parts.append(
                next_line
            )

        faculty_description = " ".join(
            description_parts
        ).strip()

    # ========================================================
    # JOIN DATE
    # ========================================================

    join_date_text = extract_value_from_same_line(
        lines,
        [
            "date of joining the institution",
            "date of joining",
            "joining date",
            "join date"
        ]
    )

    join_date = convert_join_date(
        join_date_text
    )

    # ========================================================
    # TEACHING EXPERIENCE
    # ========================================================

    teaching_experience = extract_value_from_same_line(
        lines,
        [
            "teaching experience"
        ]
    )

    # ========================================================
    # QUALIFICATIONS
    # ========================================================

    qualifications = extract_section(
        lines,
        ["qualifications"],
        [
            "teaching experience",
            "skills",
            "research & publications",
            "research and publications"
        ]
    )

    # If the PDF contains repeated Teaching Experience after
    # qualifications, the first one correctly ends the section.

    # ========================================================
    # SKILLS
    # ========================================================

    skills = extract_section(
        lines,
        ["skills"],
        [
            "research & publications",
            "research and publications",
            "journal papers",
            "conference papers",
            "phd guidance",
            "academic contributions"
        ]
    )

    # ========================================================
    # JOURNAL PAPERS
    # ========================================================

    national_journals = ""
    international_journals = ""

    journal_start = -1

    for index, line in enumerate(lines):
        if line.lower() == "journal papers":
            journal_start = index
            break

    if journal_start != -1:
        journal_lines = []

        for line in lines[journal_start + 1:]:
            if line.lower() in (
                "conference papers",
                "phd guidance",
                "academic contributions"
            ):
                break

            journal_lines.append(line)

        for line in journal_lines:
            match = re.match(
                r"^national\s*:\s*(.*)$",
                line,
                re.IGNORECASE
            )

            if match:
                national_journals = normalize_line(
                    match.group(1)
                )
                continue

            match = re.match(
                r"^international\s*:\s*(.*)$",
                line,
                re.IGNORECASE
            )

            if match:
                international_journals = normalize_line(
                    match.group(1)
                )

    # ========================================================
    # CONFERENCE PAPERS
    # ========================================================

    national_conferences = ""
    international_conferences = ""

    conference_start = -1

    for index, line in enumerate(lines):
        if line.lower() == "conference papers":
            conference_start = index
            break

    if conference_start != -1:
        conference_lines = []

        for line in lines[conference_start + 1:]:
            if line.lower() in (
                "phd guidance",
                "academic contributions"
            ):
                break

            conference_lines.append(line)

        for line in conference_lines:
            match = re.match(
                r"^national\s*:\s*(.*)$",
                line,
                re.IGNORECASE
            )

            if match:
                national_conferences = normalize_line(
                    match.group(1)
                )
                continue

            match = re.match(
                r"^international\s*:\s*(.*)$",
                line,
                re.IGNORECASE
            )

            if match:
                international_conferences = normalize_line(
                    match.group(1)
                )

    # ========================================================
    # PHD GUIDANCE
    # ========================================================

    phd_guidance = ""

    phd_start = -1

    for index, line in enumerate(lines):
        if line.lower() == "phd guidance":
            phd_start = index
            break

    if phd_start != -1:
        phd_lines = []

        for line in lines[phd_start + 1:]:
            lower_line = line.lower()

            if (
                lower_line.startswith(
                    "phds / projects guided"
                )
                or lower_line.startswith(
                    "phds/projects guided"
                )
                or lower_line == "academic contributions"
            ):
                break

            phd_lines.append(line)

        phd_guidance = normalize_line(
            " ".join(phd_lines)
        )

    # ========================================================
    # MASTER'S PROJECTS
    # ========================================================

    masters_projects = ""

    # Format 1:
    # Projects at Masters Level - 1
    for line in lines:
        match = re.search(
            r"projects\s+at\s+(?:master'?s|masters)\s+level\s*[-:]\s*(.*)$",
            line,
            re.IGNORECASE
        )

        if match:
            masters_projects = normalize_line(
                match.group(1)
            )
            break

    # Format 2:
    # PhDs / Projects Guided Projects at Masters Level - 1
    if not masters_projects:
        for line in lines:
            match = re.search(
                r"projects\s+at\s+(?:master'?s|masters)\s+level\s*[-:]\s*(.*)$",
                line,
                re.IGNORECASE
            )

            if match:
                masters_projects = normalize_line(
                    match.group(1)
                )
                break

    # ========================================================
    # ACADEMIC CONTRIBUTIONS
    # ========================================================

    academic_start = -1

    for index, line in enumerate(lines):
        if line.lower() == "academic contributions":
            academic_start = index
            break

    academic_lines = []

    if academic_start != -1:
        academic_lines = lines[academic_start + 1:]

    # ========================================================
    # BOOKS / IPR / PATENTS
    # ========================================================

    books_patents = ""

    for line in academic_lines:
        match = re.match(
            r"^books\s*/\s*ipr\s*/\s*patents\s*:\s*(.*)$",
            line,
            re.IGNORECASE
        )

        if match:
            books_patents = normalize_line(
                match.group(1)
            )
            break

    # ========================================================
    # PROFESSIONAL MEMBERSHIPS
    # ========================================================

    professional_memberships = ""

    membership_start = -1

    for index, line in enumerate(academic_lines):
        match = re.match(
            r"^professional memberships\s*:?\s*(.*)$",
            line,
            re.IGNORECASE
        )

        if match:
            membership_start = index
            first_value = normalize_line(
                match.group(1)
            )

            values = []

            if first_value:
                values.append(first_value)

            for next_line in academic_lines[index + 1:]:
                lower_next = next_line.lower()

                if (
                    lower_next.startswith(
                        "consultancy activities"
                    )
                    or lower_next.startswith("awards")
                    or lower_next.startswith("grants")
                ):
                    break

                values.append(next_line)

            professional_memberships = normalize_line(
                " ".join(values)
            )
            break

    # ========================================================
    # CONSULTANCY
    # ========================================================

    consultancy = ""

    for line in academic_lines:
        match = re.match(
            r"^consultancy activities\s*:?\s*(.*)$",
            line,
            re.IGNORECASE
        )

        if match:
            consultancy = normalize_line(
                match.group(1)
            )
            break

    # ========================================================
    # AWARDS
    # ========================================================

    awards = ""

    for index, line in enumerate(academic_lines):
        match = re.match(
            r"^awards\s*:?\s*(.*)$",
            line,
            re.IGNORECASE
        )

        if match:
            values = []

            first_value = normalize_line(
                match.group(1)
            )

            if first_value:
                values.append(first_value)

            for next_line in academic_lines[index + 1:]:
                if next_line.lower().startswith("grants"):
                    break

                values.append(next_line)

            awards = normalize_line(
                " ".join(values)
            )
            break

    # ========================================================
    # GRANTS
    # ========================================================

    grants = ""

    for line in academic_lines:
        match = re.match(
            r"^grants\s*:?\s*(.*)$",
            line,
            re.IGNORECASE
        )

        if match:
            grants = normalize_line(
                match.group(1)
            )
            break

    # ========================================================
    # FINAL DATA
    # ========================================================

    return {
        "faculty_name": faculty_name,
        "department": department,
        "designation": designation,
        "faculty_description": faculty_description,
        "join_date": join_date,
        "teaching_experience": teaching_experience,
        "qualifications": qualifications,
        "skills": skills,
        "books_patents": books_patents,
        "professional_memberships": professional_memberships,
        "consultancy": consultancy,
        "awards": awards,
        "grants": grants,
        "national_journals": national_journals,
        "international_journals": international_journals,
        "national_conferences": national_conferences,
        "international_conferences": international_conferences,
        "phd_guidance": phd_guidance,
        "masters_projects": masters_projects
    }


# ============================================================
# GET ALL FACULTY
# ============================================================

@faculty_api.route(
    "/faculties",
    methods=["GET"]
)
def get_faculties():

    connection = get_connection()

    cursor = connection.cursor(
        dictionary=True
    )

    cursor.execute("""
        SELECT
            f.faculty_id,
            f.faculty_name,
            f.department_id,
            d.department_code,
            d.department_name,
            f.designation,
            f.max_workload,
            f.status
        FROM faculty f
        JOIN department d
            ON f.department_id = d.department_id
        ORDER BY f.faculty_id
    """)

    faculties = cursor.fetchall()

    cursor.close()
    connection.close()

    return jsonify(
        faculties
    )


# ============================================================
# GET SINGLE FACULTY DETAILS
# ============================================================

@faculty_api.route(
    "/faculties/<int:faculty_id>",
    methods=["GET"]
)
def get_faculty(
    faculty_id
):

    connection = get_connection()

    cursor = connection.cursor(
        dictionary=True
    )

    cursor.execute("""
        SELECT

            f.faculty_id,
            f.faculty_name,

            f.department_id,
            d.department_code,
            d.department_name,

            f.designation,
            f.max_workload,
            f.status,

            fd.faculty_description,
            fd.join_date,
            fd.teaching_experience,

            fd.qualifications,
            fd.skills,

            fd.books_patents,
            fd.professional_memberships,
            fd.consultancy,
            fd.awards,
            fd.grants,

            fd.national_journals,
            fd.international_journals,

            fd.national_conferences,
            fd.international_conferences,

            fd.phd_guidance,
            fd.masters_projects,

            fd.profile_image

        FROM faculty f

        JOIN department d
            ON f.department_id =
               d.department_id

        LEFT JOIN faculty_details fd
            ON f.faculty_id =
               fd.faculty_id

        WHERE f.faculty_id = %s
    """, (
        faculty_id,
    ))

    faculty = cursor.fetchone()

    cursor.close()
    connection.close()

    if not faculty:

        return jsonify({
            "message":
                "Faculty not found."
        }), 404

    return jsonify(
        faculty
    )


# ============================================================
# ADD FACULTY
# ============================================================

@faculty_api.route(
    "/faculties",
    methods=["POST"]
)
def add_faculty():

    data = request.get_json()

    faculty_name = data.get(
        "faculty_name"
    )

    department = data.get(
        "department"
    )

    designation = data.get(
        "designation"
    )

    status = data.get(
        "status",
        "Active"
    )


    if (
        not faculty_name
        or not department
        or not designation
    ):

        return jsonify({
            "message":
                "Faculty name, department and designation are required."
        }), 400


    max_workload = (
        get_workload_for_designation(
            designation
        )
    )


    connection = get_connection()

    cursor = connection.cursor()


    try:

        cursor.execute("""
            SELECT department_id
            FROM department
            WHERE department_code = %s
        """, (
            department,
        ))

        department_row = cursor.fetchone()


        if not department_row:

            return jsonify({
                "message":
                    "Department not found."
            }), 404


        department_id = (
            department_row[0]
        )


        cursor.execute("""
            INSERT INTO faculty (
                faculty_name,
                department_id,
                designation,
                max_workload,
                status
            )
            VALUES (
                %s,
                %s,
                %s,
                %s,
                %s
            )
        """, (
            faculty_name,
            department_id,
            designation,
            max_workload,
            status
        ))


        faculty_id = cursor.lastrowid


        cursor.execute("""
            INSERT INTO faculty_details (
                faculty_id
            )
            VALUES (%s)
        """, (
            faculty_id,
        ))


        connection.commit()


        return jsonify({

            "message":
                "Faculty added successfully!",

            "faculty_id":
                faculty_id,

            "max_workload":
                max_workload

        }), 201


    except Exception as error:

        connection.rollback()

        print(
            "Add faculty error:",
            error
        )

        return jsonify({
            "message":
                "Unable to add faculty."
        }), 500


    finally:

        cursor.close()
        connection.close()


# ============================================================
# UPDATE FACULTY
# ============================================================

@faculty_api.route(
    "/faculties/<int:faculty_id>",
    methods=["PUT"]
)
def update_faculty(
    faculty_id
):

    data = request.get_json()

    faculty_name = data.get(
        "faculty_name"
    )

    department = data.get(
        "department"
    )

    designation = data.get(
        "designation"
    )

    status = data.get(
        "status",
        "Active"
    )


    if (
        not faculty_name
        or not department
        or not designation
    ):

        return jsonify({
            "message":
                "Faculty name, department and designation are required."
        }), 400


    max_workload = (
        get_workload_for_designation(
            designation
        )
    )


    connection = get_connection()

    cursor = connection.cursor()


    try:

        cursor.execute("""
            SELECT faculty_id
            FROM faculty
            WHERE faculty_id = %s
        """, (
            faculty_id,
        ))


        if not cursor.fetchone():

            return jsonify({
                "message":
                    "Faculty not found."
            }), 404


        cursor.execute("""
            SELECT department_id
            FROM department
            WHERE department_code = %s
        """, (
            department,
        ))


        department_row = (
            cursor.fetchone()
        )


        if not department_row:

            return jsonify({
                "message":
                    "Department not found."
            }), 404


        department_id = (
            department_row[0]
        )


        cursor.execute("""
            UPDATE faculty
            SET
                faculty_name = %s,
                department_id = %s,
                designation = %s,
                max_workload = %s,
                status = %s
            WHERE faculty_id = %s
        """, (
            faculty_name,
            department_id,
            designation,
            max_workload,
            status,
            faculty_id
        ))


        cursor.execute("""
            SELECT faculty_detail_id
            FROM faculty_details
            WHERE faculty_id = %s
        """, (
            faculty_id,
        ))


        details_exists = (
            cursor.fetchone()
        )


        if not details_exists:

            cursor.execute("""
                INSERT INTO faculty_details (
                    faculty_id
                )
                VALUES (%s)
            """, (
                faculty_id,
            ))


        connection.commit()


        return jsonify({

            "message":
                "Faculty updated successfully!",

            "max_workload":
                max_workload

        })


    except Exception as error:

        connection.rollback()

        print(
            "Update faculty error:",
            error
        )

        return jsonify({
            "message":
                "Unable to update faculty."
        }), 500


    finally:

        cursor.close()
        connection.close()


# ============================================================
# UPDATE FACULTY DETAILS
# ============================================================

@faculty_api.route(
    "/faculties/<int:faculty_id>/details",
    methods=["PUT"]
)
def update_faculty_details(
    faculty_id
):

    data = request.get_json()

    connection = get_connection()

    cursor = connection.cursor()


    try:

        cursor.execute("""
            SELECT faculty_id
            FROM faculty
            WHERE faculty_id = %s
        """, (
            faculty_id,
        ))


        if not cursor.fetchone():

            return jsonify({
                "message":
                    "Faculty not found."
            }), 404


        faculty_description = data.get(
            "faculty_description"
        )

        join_date = convert_join_date(
            data.get("join_date")
        )

        teaching_experience = data.get(
            "teaching_experience"
        )

        qualifications = data.get(
            "qualifications"
        )

        skills = data.get(
            "skills"
        )

        books_patents = data.get(
            "books_patents"
        )

        professional_memberships = data.get(
            "professional_memberships"
        )

        consultancy = data.get(
            "consultancy"
        )

        awards = data.get(
            "awards"
        )

        grants = data.get(
            "grants"
        )

        national_journals = data.get(
            "national_journals"
        )

        international_journals = data.get(
            "international_journals"
        )

        national_conferences = data.get(
            "national_conferences"
        )

        international_conferences = data.get(
            "international_conferences"
        )

        phd_guidance = data.get(
            "phd_guidance"
        )

        masters_projects = data.get(
            "masters_projects"
        )

        profile_image = data.get(
            "profile_image"
        )


        cursor.execute("""
            SELECT faculty_detail_id
            FROM faculty_details
            WHERE faculty_id = %s
        """, (
            faculty_id,
        ))


        details_row = cursor.fetchone()


        if details_row:

            cursor.execute("""
                UPDATE faculty_details
                SET

                    faculty_description = %s,
                    join_date = %s,
                    teaching_experience = %s,

                    qualifications = %s,
                    skills = %s,

                    books_patents = %s,
                    professional_memberships = %s,
                    consultancy = %s,
                    awards = %s,
                    grants = %s,

                    national_journals = %s,
                    international_journals = %s,

                    national_conferences = %s,
                    international_conferences = %s,

                    phd_guidance = %s,
                    masters_projects = %s,

                    profile_image = %s

                WHERE faculty_id = %s
            """, (

                faculty_description,
                join_date,
                teaching_experience,

                qualifications,
                skills,

                books_patents,
                professional_memberships,
                consultancy,
                awards,
                grants,

                national_journals,
                international_journals,

                national_conferences,
                international_conferences,

                phd_guidance,
                masters_projects,

                profile_image,

                faculty_id
            ))


        else:

            cursor.execute("""
                INSERT INTO faculty_details (

                    faculty_id,

                    faculty_description,
                    join_date,
                    teaching_experience,

                    qualifications,
                    skills,

                    books_patents,
                    professional_memberships,
                    consultancy,
                    awards,
                    grants,

                    national_journals,
                    international_journals,

                    national_conferences,
                    international_conferences,

                    phd_guidance,
                    masters_projects,

                    profile_image
                )

                VALUES (

                    %s,

                    %s,
                    %s,
                    %s,

                    %s,
                    %s,

                    %s,
                    %s,
                    %s,
                    %s,
                    %s,

                    %s,
                    %s,

                    %s,
                    %s,

                    %s,
                    %s,

                    %s
                )
            """, (

                faculty_id,

                faculty_description,
                join_date,
                teaching_experience,

                qualifications,
                skills,

                books_patents,
                professional_memberships,
                consultancy,
                awards,
                grants,

                national_journals,
                international_journals,

                national_conferences,
                international_conferences,

                phd_guidance,
                masters_projects,

                profile_image
            ))


        connection.commit()


        return jsonify({
            "message":
                "Faculty details updated successfully!"
        })


    except Exception as error:

        connection.rollback()

        print(
            "Faculty details update error:",
            error
        )

        return jsonify({
            "message":
                "Unable to update faculty details."
        }), 500


    finally:

        cursor.close()
        connection.close()


# ============================================================
# IMPORT FACULTY
# ============================================================

@faculty_api.route(
    "/faculties/import",
    methods=["POST"]
)
def import_faculty():

    if "file" not in request.files:

        return jsonify({
            "message":
                "No faculty file selected."
        }), 400


    file = request.files["file"]


    if not file or not file.filename:

        return jsonify({
            "message":
                "Invalid faculty file."
        }), 400


    original_filename = (
        file.filename
    )


    extension = os.path.splitext(
        original_filename
    )[1].lower()


    if (
        extension
        not in ALLOWED_FACULTY_EXTENSIONS
    ):

        return jsonify({
            "message":
                "Unsupported faculty file type."
        }), 400


    filename = secure_filename(
        original_filename
    )


    file_path = os.path.join(
        FACULTY_UPLOAD_FOLDER,
        filename
    )


    try:

        file.save(
            file_path
        )


        # ----------------------------------------------------
        # Extract
        # ----------------------------------------------------

        extracted_text = (
            extract_faculty_text(
                file_path,
                extension
            )
        )


        if not extracted_text.strip():

            return jsonify({
                "message":
                    "Unable to extract faculty information from the file."
            }), 400


        print(
            "\n========== FACULTY IMPORT TEXT =========="
        )

        print(
            extracted_text[:10000]
        )

        print(
            "=========================================\n"
        )


        # ----------------------------------------------------
        # Parse
        # ----------------------------------------------------

        faculty_data = (
            parse_faculty_information(
                extracted_text
            )
        )


        faculty_name = (
            faculty_data[
                "faculty_name"
            ]
        )

        department = (
            faculty_data[
                "department"
            ]
        )

        designation = (
            faculty_data[
                "designation"
            ]
        )


        # ----------------------------------------------------
        # Debug information
        # ----------------------------------------------------

        print(
            "Detected Faculty Name:",
            faculty_name
        )

        print(
            "Detected Department:",
            department
        )

        print(
            "Detected Designation:",
            designation
        )


        # ----------------------------------------------------
        # Required fields
        # ----------------------------------------------------

        if not faculty_name:

            return jsonify({
                "message":
                    "Faculty name could not be detected."
            }), 400


        if not designation:

            return jsonify({
                "message":
                    "Faculty designation could not be detected."
            }), 400


        if not department:

            return jsonify({
                "message":
                    "Faculty department could not be detected."
            }), 400


        # ----------------------------------------------------
        # Database
        # ----------------------------------------------------

        connection = get_connection()

        cursor = connection.cursor()


        try:

            # ------------------------------------------------
            # Find department
            # ------------------------------------------------

            cursor.execute("""
                SELECT
                    department_id,
                    department_code
                FROM department
                WHERE
                    UPPER(department_code)
                    = UPPER(%s)
                    OR LOWER(department_name)
                    = LOWER(%s)
            """, (
                department,
                department
            ))


            department_row = (
                cursor.fetchone()
            )


            if not department_row:

                return jsonify({

                    "message":
                        f"Department '{department}' was not found."

                }), 404


            department_id = (
                department_row[0]
            )

            database_department_code = (
                department_row[1]
            )


            # ------------------------------------------------
            # Workload
            # ------------------------------------------------

            max_workload = (
                get_workload_for_designation(
                    designation
                )
            )


            # ------------------------------------------------
            # Duplicate
            # ------------------------------------------------

            cursor.execute("""
                SELECT
                    faculty_id
                FROM faculty
                WHERE
                    LOWER(faculty_name)
                    = LOWER(%s)
                    AND department_id = %s
            """, (
                faculty_name,
                department_id
            ))


            existing = (
                cursor.fetchone()
            )


            if existing:

                return jsonify({

                    "message":
                        "Faculty already exists in this department.",

                    "faculty_id":
                        existing[0]

                }), 409


            # ------------------------------------------------
            # Insert main faculty
            # ------------------------------------------------

            cursor.execute("""
                INSERT INTO faculty (
                    faculty_name,
                    department_id,
                    designation,
                    max_workload,
                    status
                )
                VALUES (
                    %s,
                    %s,
                    %s,
                    %s,
                    %s
                )
            """, (
                faculty_name,
                department_id,
                designation,
                max_workload,
                "Active"
            ))


            faculty_id = (
                cursor.lastrowid
            )


            # ------------------------------------------------
            # Profile image
            # ------------------------------------------------

            profile_image = None


            if extension in (
                ".png",
                ".jpg",
                ".jpeg",
                ".webp"
            ):

                profile_image = (
                    file_path.replace(
                        "\\",
                        "/"
                    )
                )


            # ------------------------------------------------
            # Insert details
            # ------------------------------------------------

            cursor.execute("""
                INSERT INTO faculty_details (

                    faculty_id,

                    faculty_description,
                    join_date,
                    teaching_experience,

                    qualifications,
                    skills,

                    books_patents,
                    professional_memberships,
                    consultancy,
                    awards,
                    grants,

                    national_journals,
                    international_journals,

                    national_conferences,
                    international_conferences,

                    phd_guidance,
                    masters_projects,

                    profile_image
                )

                VALUES (

                    %s,

                    %s,
                    %s,
                    %s,

                    %s,
                    %s,

                    %s,
                    %s,
                    %s,
                    %s,
                    %s,

                    %s,
                    %s,

                    %s,
                    %s,

                    %s,
                    %s,

                    %s
                )
            """, (

                faculty_id,

                faculty_data[
                    "faculty_description"
                ],

                convert_join_date(
                    faculty_data[
                        "join_date"
                    ]
                ),

                faculty_data[
                    "teaching_experience"
                ],

                faculty_data[
                    "qualifications"
                ],

                faculty_data[
                    "skills"
                ],

                faculty_data[
                    "books_patents"
                ],

                faculty_data[
                    "professional_memberships"
                ],

                faculty_data[
                    "consultancy"
                ],

                faculty_data[
                    "awards"
                ],

                faculty_data[
                    "grants"
                ],

                faculty_data[
                    "national_journals"
                ],

                faculty_data[
                    "international_journals"
                ],

                faculty_data[
                    "national_conferences"
                ],

                faculty_data[
                    "international_conferences"
                ],

                faculty_data[
                    "phd_guidance"
                ],

                faculty_data[
                    "masters_projects"
                ],

                profile_image
            ))


            connection.commit()


            print(
                "Faculty imported successfully:"
            )

            print(
                "Faculty ID:",
                faculty_id
            )

            print(
                "Name:",
                faculty_name
            )

            print(
                "Department:",
                database_department_code
            )

            print(
                "Designation:",
                designation
            )

            print(
                "Workload:",
                max_workload
            )


            return jsonify({

                "message":
                    "Faculty imported successfully!",

                "faculty_id":
                    faculty_id,

                "faculty_name":
                    faculty_name,

                "department":
                    database_department_code,

                "designation":
                    designation,

                "max_workload":
                    max_workload

            }), 201


        except Exception as error:

            connection.rollback()

            print(
                "Faculty import database error:",
                error
            )

            return jsonify({

                "message":
                    "Unable to import faculty.",

                "error":
                    str(error)

            }), 500


        finally:

            cursor.close()
            connection.close()


    except Exception as error:

        print(
            "Faculty import error:",
            error
        )

        return jsonify({

            "message":
                "Unable to process faculty file.",

            "error":
                str(error)

        }), 500

# ============================================================
# UPLOAD FACULTY PROFILE IMAGE
# ============================================================

@faculty_api.route(
    "/faculties/<int:faculty_id>/profile-image",
    methods=["POST"]
)
def upload_faculty_profile_image(faculty_id):

    # --------------------------------------------------------
    # Check file
    # --------------------------------------------------------

    if "profile_image" not in request.files:

        return jsonify({
            "message":
                "No profile image selected."
        }), 400


    file = request.files["profile_image"]


    if not file or not file.filename:

        return jsonify({
            "message":
                "Invalid profile image."
        }), 400


    # --------------------------------------------------------
    # Check extension
    # --------------------------------------------------------

    extension = os.path.splitext(
        file.filename
    )[1].lower()


    allowed_extensions = {
        ".jpg",
        ".jpeg",
        ".png",
        ".webp"
    }


    if extension not in allowed_extensions:

        return jsonify({
            "message":
                "Only JPG, JPEG, PNG and WEBP images are allowed."
        }), 400


    # --------------------------------------------------------
    # Check faculty exists
    # --------------------------------------------------------

    connection = get_connection()

    cursor = connection.cursor()


    try:

        cursor.execute("""
            SELECT faculty_id
            FROM faculty
            WHERE faculty_id = %s
        """, (
            faculty_id,
        ))


        faculty = cursor.fetchone()


        if not faculty:

            return jsonify({
                "message":
                    "Faculty not found."
            }), 404


        # ----------------------------------------------------
        # Create faculty image folder
        # ----------------------------------------------------

        faculty_folder = os.path.join(
            FACULTY_UPLOAD_FOLDER,
            "profile_images"
        )


        os.makedirs(
            faculty_folder,
            exist_ok=True
        )


        # ----------------------------------------------------
        # Generate safe filename
        # ----------------------------------------------------

        original_name = secure_filename(
            file.filename
        )


        name_without_extension = os.path.splitext(
            original_name
        )[0]


        # Remove unsafe characters
        name_without_extension = re.sub(
            r"[^A-Za-z0-9_-]",
            "_",
            name_without_extension
        )


        filename = (
            f"{faculty_id}_"
            f"{name_without_extension}"
            f"{extension}"
        )


        file_path = os.path.join(
            faculty_folder,
            filename
        )


        # ----------------------------------------------------
        # Delete old profile images
        # ----------------------------------------------------

        for existing_file in os.listdir(
            faculty_folder
        ):

            if existing_file.startswith(
                f"{faculty_id}_"
            ):

                old_file_path = os.path.join(
                    faculty_folder,
                    existing_file
                )

                try:

                    if os.path.isfile(
                        old_file_path
                    ):

                        os.remove(
                            old_file_path
                        )

                except Exception as error:

                    print(
                        "Unable to remove old image:",
                        error
                    )


        # ----------------------------------------------------
        # Save new image
        # ----------------------------------------------------

        file.save(
            file_path
        )


        # ----------------------------------------------------
        # Path stored in database
        # ----------------------------------------------------

        database_path = (
            f"uploads/faculty/"
            f"profile_images/"
            f"{filename}"
        )


        # ----------------------------------------------------
        # Check faculty_details row
        # ----------------------------------------------------

        cursor.execute("""
            SELECT faculty_detail_id
            FROM faculty_details
            WHERE faculty_id = %s
        """, (
            faculty_id,
        ))


        details = cursor.fetchone()


        if details:

            cursor.execute("""
                UPDATE faculty_details
                SET profile_image = %s
                WHERE faculty_id = %s
            """, (
                database_path,
                faculty_id
            ))

        else:

            cursor.execute("""
                INSERT INTO faculty_details (
                    faculty_id,
                    profile_image
                )
                VALUES (
                    %s,
                    %s
                )
            """, (
                faculty_id,
                database_path
            ))


        connection.commit()


        return jsonify({

            "message":
                "Faculty profile image uploaded successfully.",

            "faculty_id":
                faculty_id,

            "profile_image":
                database_path

        }), 200


    except Exception as error:

        connection.rollback()

        print(
            "Faculty profile image upload error:",
            error
        )

        return jsonify({

            "message":
                "Unable to upload faculty profile image.",

            "error":
                str(error)

        }), 500


    finally:

        cursor.close()
        connection.close()
# ============================================================
# SERVE FACULTY PROFILE IMAGES
# ============================================================

@faculty_api.route(
    "/uploads/faculty/profile_images/<path:filename>",
    methods=["GET"]
)
def serve_faculty_profile_image(filename):

    profile_image_folder = os.path.join(
        FACULTY_UPLOAD_FOLDER,
        "profile_images"
    )

    return send_from_directory(
        profile_image_folder,
        filename
    )
# ============================================================
# DELETE FACULTY
# ============================================================

@faculty_api.route(
    "/faculties/<int:faculty_id>",
    methods=["DELETE"]
)
def delete_faculty(
    faculty_id
):

    connection = get_connection()

    cursor = connection.cursor()


    try:

        cursor.execute("""
            SELECT faculty_id
            FROM faculty
            WHERE faculty_id = %s
        """, (
            faculty_id,
        ))


        if not cursor.fetchone():

            return jsonify({
                "message":
                    "Faculty not found."
            }), 404


        cursor.execute("""
            DELETE FROM faculty
            WHERE faculty_id = %s
        """, (
            faculty_id,
        ))


        connection.commit()


        return jsonify({
            "message":
                "Faculty deleted successfully!"
        })


    except Exception as error:

        connection.rollback()

        print(
            "Delete faculty error:",
            error
        )

        return jsonify({
            "message":
                "Unable to delete faculty."
        }), 500


    finally:

        cursor.close()
        connection.close()